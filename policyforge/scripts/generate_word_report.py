#!/usr/bin/env python3
"""
Generate PolicyForge Word Report (.docx)

Structure:
- Two-page body + bibliography page
- Abstract, Industry Relevance, Methodology, Results, Implications, Conclusion
- APA format citations
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx

def add_page_break(doc):
    doc.add_page_break()

def set_font(run, bold=False, size=12, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, size=14, bold=True, color=None, spacing_before=12, spacing_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_body(doc, text, size=11, spacing_after=6, indent=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def build_report():
    doc = Document()

    # ── Page margins ─────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

    # ── TITLE BLOCK ──────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    t = title_p.add_run("PolicyForge: Automated Healthcare Policy Extraction")
    t.bold = True
    t.font.size = Pt(15)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(4)
    s = subtitle_p.add_run(
        "Content Management in Health Care: Converting Billing Policies into\n"
        "Executable Rules Using Large Language Models"
    )
    s.font.size = Pt(11)
    s.italic = True

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(10)
    m = meta_p.add_run("Shristi Kumar  |  Cotiviti Intern Assessment  |  July 2026")
    m.font.size = Pt(10)
    m.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # horizontal rule (thin border paragraph)
    hr = doc.add_paragraph()
    hr.paragraph_format.space_after = Pt(8)
    hr_run = hr.add_run("─" * 80)
    hr_run.font.size = Pt(8)
    hr_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── ABSTRACT ─────────────────────────────────────────────
    add_heading(doc, "Abstract", size=13, color=(0x1F, 0x49, 0x7D),
                spacing_before=4, spacing_after=4)

    add_body(doc,
        "This report presents PolicyForge, a proof-of-concept system that automates extraction "
        "of Medicare billing rules from National Coverage Determinations (NCDs) and Code of Federal "
        "Regulations (CFRs). The system uses a six-agent LLM pipeline with hybrid retrieval "
        "(BM25 + FAISS) to convert unstructured policy text into structured HCPCS codes and "
        "frequency limits, then applies those rules to flag high-utilization providers in CMS "
        "Part B data. We evaluated the system on 15 real CMS policies with hand-labeled gold "
        "standards, achieving 98.2% mean F1 on code extraction and 96.4% weighted F1 when "
        "scored by clinical severity. The project demonstrates that LLM-based policy extraction "
        "can reduce analyst time from 45 minutes to 8 seconds per policy (15× cost reduction), "
        "but also shows why human review remains necessary for high-risk cancer-screening policies.",
        spacing_after=6)

    # ── SECTION 1: INDUSTRY RELEVANCE ───────────────────────
    add_heading(doc, "1. Industry Relevance", size=13, color=(0x1F, 0x49, 0x7D),
                spacing_before=4, spacing_after=4)

    add_body(doc,
        "For companies like Cotiviti that support payer audit and payment integrity programs, "
        "healthcare content management is a core operational bottleneck. Every Medicare policy "
        "update must be translated into executable logic: which HCPCS/CPT codes are covered, "
        "how often services may be billed, and which patient populations qualify. Today this "
        "work is done manually by medical policy analysts at roughly 45 minutes and $56 per "
        "policy. With CMS issuing over 1,200 coverage updates annually, manual extraction does "
        "not scale and introduces inconsistency across audit teams.",
        spacing_after=5)

    add_body(doc,
        "The industry need is not simply faster document reading. Payers need repeatable, "
        "auditable rules that can be applied to claims data at scale. PolicyForge addresses "
        "this by connecting policy ingestion directly to provider-level utilization analysis, "
        "so extracted rules can immediately support audit triage rather than sitting in static "
        "spreadsheets.",
        spacing_after=5)

    # ── SECTION 2: METHODOLOGY ───────────────────────────────
    add_heading(doc, "2. Methodology", size=13, color=(0x1F, 0x49, 0x7D),
                spacing_before=4, spacing_after=4)

    add_body(doc,
        "We built PolicyForge as an end-to-end pipeline with four stages:",
        spacing_after=3)

    add_bullet(doc,
        "Data collection. We sourced 15 real Medicare policies from CMS.gov and eCFR, spanning "
        "cancer screening (colorectal, mammography, lung), cardiovascular and metabolic "
        "screening, and behavioral health programs.")
    add_bullet(doc,
        "Policy extraction. A LangGraph orchestration runs six agents in sequence: Retriever "
        "(hybrid RAG over policy text), Extractor (Mistral-large with structured JSON output), "
        "Critic (completeness check), Compiler (policy-to-code translation), Adjudicator "
        "(statistical outlier detection), and Explainer (audit memo generation).")
    add_bullet(doc,
        "Evaluation. We created manual gold standards by reading each policy document and "
        "recording expected HCPCS codes and frequency limits. LLM outputs were scored with "
        "precision, recall, and F1 against these gold standards. We also classified policies "
        "by clinical severity (Tier 1: cancer; Tier 2: CVD/metabolic; Tier 3: routine) to "
        "compute a weighted F1 that reflects patient-harm risk.")
    add_bullet(doc,
        "Utilization analysis. Extracted rules were applied to a CMS Part B provider summary "
        "dataset (21,521 providers). Providers exceeding mean + 2 standard deviations on "
        "services-per-beneficiary were flagged for audit review.")

    add_body(doc,
        "To improve extraction quality, we iterated on three technical levers: extending LLM "
        "context from 4K to 12K characters, adding few-shot examples of correct extractions, "
        "and splitting extraction into separate passes for codes and frequency limits.",
        spacing_after=5)

    # ── PAGE BREAK ───────────────────────────────────────────
    add_page_break(doc)

    # ── SECTION 3: RESULTS & OUTPUT ──────────────────────────
    add_heading(doc, "3. Results and Output", size=13, color=(0x1F, 0x49, 0x7D),
                spacing_before=4, spacing_after=4)

    add_body(doc,
        "Table 1 summarizes the primary outputs of the proof of concept.",
        spacing_after=3)

    add_bullet(doc,
        "Extraction accuracy: 98.2% mean HCPCS F1 across 15 policies; 14 of 15 policies at F1 ≥ 0.9.")
    add_bullet(doc,
        "Clinical-weighted accuracy: 96.4% weighted F1 when cancer-screening policies receive "
        "higher weight than routine behavioral health policies.")
    add_bullet(doc,
        "Operational efficiency: 8 seconds and $3.75 per policy (including review time) vs. "
        "45 minutes and $56.25 for fully manual extraction.")
    add_bullet(doc,
        "Audit triage output: 389 of 21,521 providers (1.8%) flagged as statistical outliers, "
        "replacing an initial broken baseline that flagged 100% of providers.")
    add_bullet(doc,
        "Structured deliverables: JSON extraction files per policy, evaluation reports, and "
        "a reproducible GitHub repository with source code and gold standards.")

    add_body(doc,
        "Performance was not uniform across all policy types. Tier 2 and Tier 3 policies "
        "(cardiovascular, diabetes, depression, obesity) reached 100% F1. Tier 1 cancer "
        "screening averaged 93.3%, with NCD 210.3 (colorectal screening) remaining at 80% F1 "
        "because two of eleven HCPCS codes were still missed after iterative improvement. "
        "This gap matters clinically: missing colonoscopy codes could lead to incorrect claim "
        "denials and delayed screening.",
        spacing_after=5)

    add_heading(doc, "4. Implications for Cotiviti", size=13, color=(0x1F, 0x49, 0x7D),
                spacing_before=4, spacing_after=4)

    add_body(doc,
        "The results support a staged adoption path rather than a single go/no-go decision on "
        "full automation. In the near term, PolicyForge is best positioned as a policy "
        "ingestion and audit-prioritization tool: analysts use LLM extraction to draft rules "
        "quickly, then review flagged providers rather than reviewing the entire provider "
        "population. This preserves human oversight where clinical risk is highest while "
        "still capturing most of the 15× efficiency gain.",
        spacing_after=5)

    add_body(doc,
        "Broader industry trends reinforce this approach. LLM accuracy on regulatory text is "
        "now strong enough for triage but not yet sufficient for unsupervised adjudication on "
        "life-critical policies. Regulatory requirements (CMS audit trails, potential FDA "
        "review of clinical decision support tools) also favor human-in-the-loop deployment. "
        "The immediate opportunity for Cotiviti is to reduce policy analyst workload and "
        "narrow audit scope to high-risk providers. The longer-term opportunity is hybrid "
        "automation on lower-risk policy tiers once independent validation (NCCI cross-check, "
        "second medical coder review) confirms extraction quality.",
        spacing_after=5)

    add_body(doc,
        "Key risks to manage include self-validation bias (gold standards created by the same "
        "team that built the system), residual errors on complex multi-code policies like "
        "colorectal screening, and the need for confidence scoring so uncertain extractions "
        "route automatically to expert review.",
        spacing_after=5)

    # ── CONCLUSION ────────────────────────────────────────────
    add_heading(doc, "5. Conclusion", size=13, color=(0x1F, 0x49, 0x7D),
                spacing_before=4, spacing_after=4)
    add_body(doc,
        "PolicyForge shows that multi-agent LLM orchestration can convert Medicare policy "
        "documents into structured, data-ready rules at scale. The proof of concept is "
        "technically sound (98.2% mean F1), operationally relevant (1.8% targeted audit "
        "rate), and economically justified (15× cost reduction). The honest limitation is "
        "clinical: high-stakes screening policies still require human validation before "
        "any automated denial logic. For Cotiviti, the strategic value is using AI to "
        "accelerate content management and focus expert reviewers where patient impact is "
        "greatest, not replacing those reviewers entirely.",
        spacing_after=6)

    # ── PAGE BREAK to bibliography ────────────────────────────
    add_page_break(doc)

    # ── BIBLIOGRAPHY (APA) ────────────────────────────────────
    add_heading(doc, "References", size=13, color=(0x1F, 0x49, 0x7D),
                spacing_before=4, spacing_after=8)

    refs = [
        ("Centers for Medicare & Medicaid Services. (2023). ",
         "National coverage determinations (NCD) manual (Pub. 100-3). ",
         "U.S. Department of Health and Human Services. https://www.cms.gov/medicare-coverage-database/"),

        ("Centers for Medicare & Medicaid Services. (2024). ",
         "National Correct Coding Initiative (NCCI) methodology. ",
         "https://www.cms.gov/medicare/coding-billing/national-correct-coding-initiative-ncci"),

        ("Code of Federal Regulations, 42 C.F.R. § 410.18. (2023). ",
         "Diabetes screening tests. ",
         "U.S. Government Publishing Office. https://www.ecfr.gov/"),

        ("Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., Küttler, H., Lewis, M., Yih, W., Rocktäschel, T., Riedel, S., & Kiela, D. (2020). ",
         "Retrieval-augmented generation for knowledge-intensive NLP tasks. ",
         "Advances in Neural Information Processing Systems, 33, 9459–9474. https://arxiv.org/abs/2005.11401"),

        ("Mistral AI. (2024). ",
         "Mistral large model documentation. ",
         "https://docs.mistral.ai/getting-started/models/"),

        ("Hong, S., Zhuge, M., Chen, J., Zheng, X., Cheng, Y., Wang, J., Zhang, C., Wang, Z., Yau, S. K. S., Lin, Z., Zhou, L., Ran, C., Xiao, L., Wu, C., & Schmidhuber, J. (2024). ",
         "MetaGPT: Meta programming for a multi-agent collaborative framework. ",
         "International Conference on Learning Representations (ICLR 2024). https://arxiv.org/abs/2308.00352"),
    ]

    for author_part, title_part, source_part in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(0)

        r1 = p.add_run(author_part)
        r1.font.size = Pt(10)

        r2 = p.add_run(title_part)
        r2.font.size = Pt(10)
        r2.italic = True

        r3 = p.add_run(source_part)
        r3.font.size = Pt(10)

    # ── SAVE ──────────────────────────────────────────────────
    path = "../PolicyForge_Report.docx"
    doc.save(path)
    print(f"✅ Word report saved: {path}")
    return path


# Monkey-patch a helper used above
_orig_add_body = add_body
def add_body(doc, text, size=11, spacing_after=6, indent=False, italic=False, bold_=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    if bold_:
        run.bold = True
    return p

if __name__ == "__main__":
    build_report()
